from validation import *
from docx import Document
from datetime import date
import tkinter as tk , tkinter.messagebox, access_db

#Class to manage all people details; staff, customer
class ManagePeople:
    def __init__(self,container):
        self.container = container

    #Get highlighted row in the listbox
    def get_selected_row(self,event):
        #selected tuple is the selected item from listbox
        global selected_tuple
        index = self.container.lst_box.curselection()[0]
        selected_tuple = self.container.lst_box.get(index)
        
        #populating the fields with data upon selection
        self.container.name.delete(0,tk.END)
        self.container.name.insert(tk.END,selected_tuple[1])
        
        self.container.age.delete(0,tk.END)
        self.container.age.insert(tk.END,selected_tuple[2])
        
        self.container.phone.delete(0,tk.END)
        self.container.phone.insert(tk.END,selected_tuple[3])
        
        self.container.address.delete(0,tk.END)
        self.container.address.insert(tk.END,selected_tuple[4])

        self.container.email.delete(0,tk.END)
        self.container.email.insert(tk.END,selected_tuple[5])

        self.container.password.delete(0,tk.END)
        self.container.password.insert(tk.END,selected_tuple[6])
    
    #view all staff, customer data from the database
    def view_data(self,db):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.view(db):
            self.container.lst_box.insert(tk.END,row)

    #search data
    def search_data(self,db):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.search(self.container.name.get(),self.container.age.get(),self.container.phone.get(),self.container.address.get(),self.container.email.get(),db):
            self.container.lst_box.insert(tk.END,row)

    #insert staff,customer to the database
    def insert_data(self,db):
        #check for validation; presence check and length check
        if Validation(self.container.name.get()).presence_check() == True and Validation(self.container.phone.get()).presence_check() == True and Validation(self.container.email.get()).presence_check() == True and Validation(self.container.password.get()).presence_check() == True:
            if Validation(self.container.phone.get()).length_check(11) == True and Validation(self.container.password.get()).range_check(7) == True:
                access_db.insert(self.container.name.get(),self.container.age.get(),self.container.phone.get(),self.container.address.get(),self.container.email.get(),self.container.password.get(),db)
                self.container.lst_box.delete(0,tk.END)
                self.container.lst_box.insert(tk.END,(self.container.name.get(),self.container.age.get(),self.container.phone.get(),self.container.address.get(),self.container.email.get(),self.container.password.get()))
            
            #error box shown if data not valid
            else:
                tkinter.messagebox.showerror("Error",'Phone number must be 11 digits\n Password must be longer than 7 characters')
        else:
            tkinter.messagebox.showerror("Error",'Ensure all fields are entered')

    #delete records
    def delete_data(self,db):
        #clearing the lisbox
        self.container.lst_box.delete(0,tk.END)
        #clearing the entry widgets
        self.container.name.delete(0,tk.END)
        self.container.age.delete(0,tk.END)
        self.container.phone.delete(0,tk.END)
        self.container.address.delete(0,tk.END)
        self.container.email.delete(0,tk.END)
        self.container.password.delete(0,tk.END)

        access_db.delete(selected_tuple[0],db)

    #update records
    def update_data(self,db):
        if Validation(self.container.name.get()).presence_check() == True and Validation(self.container.phone.get()).presence_check() == True and Validation(self.container.email.get()).presence_check() == True and Validation(self.container.password.get()).presence_check() == True:
            if Validation(self.container.phone.get()).length_check(11) == True and Validation(self.container.password.get()).range_check(7) == True:
                access_db.update(selected_tuple[0],self.container.name.get(),self.container.age.get(),self.container.phone.get(),self.container.address.get(),self.container.email.get(),self.container.password.get(),db)
            
            #error box shown if data not valid
            else:
                tkinter.messagebox.showerror("Error",'Phone number must be 11 digits\n Password must be longer than 7 characters')
        else:
            tkinter.messagebox.showerror("Error",'Ensure all fields are entered')


    #sort data by given factor
    def sort(self,db,factor):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.sort(db,factor):
            self.container.lst_box.insert(tk.END,row)

#Class to manage all product details
class ManageProduct:
    def __init__(self,container):
        self.container = container
    
    #highlighted data item from the listbox is selected
    def get_selected_row(self,event):
        global selected_tuple
        index = self.container.lst_box.curselection()[0]
        selected_tuple = self.container.lst_box.get(index)
        
        #populating the fields upon selection
        self.container.name.delete(0,tk.END)
        self.container.name.insert(tk.END,selected_tuple[1])

        self.container.price.delete(0,tk.END)
        self.container.price.insert(tk.END,selected_tuple[2])
        try:
            self.container.isbn.delete(0,tk.END)
            self.container.isbn.insert(tk.END,selected_tuple[3])
        except:
            self.container.address.delete(0,tk.END)
            self.container.address.insert(tk.END,selected_tuple[4])

    #view all data from the database
    def view_data(self,db):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.view_pd(db):
            self.container.lst_box.insert(tk.END,row)

    #search for a record in the database
    def search_data(self,db):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.search_pd(self.container.name.get(),self.container.price.get(),self.container.isbn.get(),self.container.address.get(),self.container.category.get(),db):
            self.container.lst_box.insert(tk.END,row)

    #insert new products into the database
    def insert_data(self,db):
        if Validation(self.container.name.get()).presence_check() == True and Validation(self.container.price.get()).presence_check() == True and Validation(self.container.isbn.get()).presence_check() == True and Validation(self.container.category.get()).presence_check() == True:
            if Validation(self.container.isbn.get()).length_check(10) == True:
                access_db.insert_pd(self.container.name.get(),self.container.price.get(),self.container.isbn.get(),self.container.address.get(),self.container.category.get(),db)
                self.container.lst_box.delete(0,tk.END)
                self.container.lst_box.insert(tk.END,(self.container.name.get(),self.container.price.get(),self.container.isbn.get(),self.container.address.get(),self.container.category.get()))
            else:
                tkinter.messagebox.showerror("Error",'Ensure isbn is 10 digits long')
        else:
            tkinter.messagebox.showerror("Error",'Ensure all fields are entered')

    #delete selected data from the database
    def delete_data(self,db):
        self.container.lst_box.delete(0,tk.END)
        access_db.delete_pd(selected_tuple[0],db)
    
    #update records in the database
    def update_data(self,db):
        if Validation(self.container.name.get()).presence_check() == True and Validation(self.container.price.get()).presence_check() == True and Validation(self.container.isbn.get()).presence_check() == True and Validation(self.container.category.get()).presence_check() == True:
            if Validation(self.container.isbn.get()).length_check(10) == True:
                access_db.update_pd(selected_tuple[0],self.container.name.get(),self.container.price.get(),self.container.isbn.get(),self.container.address.get(),self.container.category.get(),db)
            else:
                tkinter.messagebox.showerror("Error",'Ensure isbn is 10 digits long')
        else:
            tkinter.messagebox.showerror("Error",'Ensure all fields are entered')

#Class to manage orders from customers side
class CustomerProductView:
    def __init__(self,container):
        self.container = container
    
    #searching for products
    def search_data(self,db):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.search_pd(self.container.name.get(),self.container.price.get(),self.container.isbn.get(),"##","##",db):
            self.container.lst_box.insert(tk.END,row)

    #inserting records
    def insert_order(self,user_id,product_type):
        product_id = selected_tuple[0]
        self.container.lst_box.delete(0,tk.END)
        access_db.create_order(user_id,selected_tuple[0],"Orders")

    #view the ordered products; customers basket
    def view_order(self,userid):
        total_price = 0
        self.lst_box.delete(0,tk.END)
        for row in access_db.view_order(userid):
            self.lst_box.insert(tk.END,row)
            total_price += row[1]
        
        #calculating total price and returning it
        return total_price

    #view all customer orders
    def view_all_orders(self):
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.view_all_orders():
            self.container.lst_box.insert(tk.END,row)

#Class to manage orders from staff's side
class ManageOrder:
    def __init__(self,container):
        self.container = container

    #Get data from selected row
    def get_selected_row(self,event):
        global selected_tuple
        index = self.container.lst_box.curselection()[0]
        selected_tuple = self.container.lst_box.get(index)
        
        #populating the fields upon selection for users who made order
        try:
            self.container.userid.delete(0,tk.END)
            self.container.userid.insert(tk.END,selected_tuple[0])

            self.container.cust_name.delete(0,tk.END)
            self.container.cust_name.insert(tk.END,selected_tuple[1])

            self.container.cust_address.delete(0,tk.END)
            self.container.cust_address.insert(tk.END,selected_tuple[2])

            self.container.cust_email.delete(0,tk.END)
            self.container.cust_email.insert(tk.END,selected_tuple[3])
        
        #populating the fields upon selection for the products ordered
        except:
            self.container.prod_id.delete(0,tk.END)
            self.container.prod_id.insert(tk.END,selected_tuple[0])

            self.container.prod_name.delete(0,tk.END)
            self.container.prod_name.insert(tk.END,selected_tuple[1])

            self.container.prod_price.delete(0,tk.END)
            self.container.prod_price.insert(tk.END,selected_tuple[2])

            self.container.prod_isbn.delete(0,tk.END)
            self.container.prod_isbn.insert(tk.END,selected_tuple[3])

    #Search a given order by customer id
    def search_order(self):
        userid = self.container.userid.get()
        self.container.lst_box.delete(0,tk.END)
        for row in access_db.search_order(userid):
            self.container.lst_box.insert(tk.END,row)
        
    #Display all the items in a specific user order
    def display_items(self):
        total_price = 0
        all_items = access_db.view_ordered_items(selected_tuple[0])
        self.container.lst_box.delete(0,tk.END)
        for row in all_items:
            self.container.lst_box.insert(tk.END,row)
            total_price += row[2]
        
        #calculating and returning total price
        return total_price

    #Transfering data from one table to anohter when marked complete
    def mark_complete(self):
        #clearing the lisbox
        self.container.lst_box.delete(0,tk.END)
        #clearing the entry widgets
        self.container.userid.delete(0,tk.END)
        self.container.cust_name.delete(0,tk.END)
        self.container.cust_address.delete(0,tk.END)
        self.container.cust_email.delete(0,tk.END)

        access_db.transfer_orders(selected_tuple[0])

    #download the invoice of an order 
    def download_invoice(self):
        #open the invoice template and add data specific to the user
        document = Document('res/invoice.docx')
        records = tuple(access_db.view_ordered_items(selected_tuple[0]))

        #adding texts
        document.add_paragraph(f'Date: {date.today()} ')
        document.add_paragraph(f'User ID: {selected_tuple[0]} ')
        document.add_paragraph(f'Name: {selected_tuple[1]} ')
        document.add_paragraph(f'Address: {selected_tuple[2]} ')
        document.add_paragraph(f'Total Price: {round(self.display_items(),2)}')
        document.add_paragraph('Ordered Items')

        self.container.lst_box.delete(0,tk.END)
        
        #creating table with all ordered items and their prices
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product ID'
        hdr_cells[1].text = 'Product Name'
        hdr_cells[2].text = 'Price (£)'
        for id, name, price,isbn in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = name
            row_cells[2].text = str(price)

        #invoice file creaed in formate useridnameinvoice.docx
        document.save(f'user{selected_tuple[0]}invoice.docx')
